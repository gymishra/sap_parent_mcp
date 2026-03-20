"""Convert GUIDE.md to a Word document."""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Style setup
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
title = doc.add_heading('SAP S/4HANA + Amazon Bedrock AgentCore + Okta OIDC', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('Integration Guide')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph('')

# ── Overview ──
doc.add_heading('Overview', level=1)
doc.add_paragraph(
    'This guide documents the end-to-end integration of SAP S/4HANA OData APIs '
    'with Amazon Bedrock AgentCore Runtime using Okta as the OIDC identity provider. '
    'An MCP (Model Context Protocol) server hosted on AgentCore Runtime exposes SAP '
    'OData services as tools that can be consumed by AI agents or MCP clients.'
)

doc.add_heading('Architecture Flow', level=2)
doc.add_paragraph(
    '1. Python Client authenticates with Okta (Authorization Code flow)\n'
    '2. Client receives access_token (JWT)\n'
    '3. Client sends Bearer token to AgentCore Runtime\n'
    '4. AgentCore validates JWT against Okta OIDC discovery URL\n'
    '5. MCP Server receives the token as a tool parameter\n'
    '6. MCP Server calls SAP OData API with Bearer token\n'
    '7. SAP validates token against Okta JWKS keys\n'
    '8. Sales order data flows back to the client'
)

doc.add_heading('Token Flow', level=2)
doc.add_paragraph(
    'The same Okta access_token is used for both:\n'
    '• Inbound auth: AgentCore Runtime validates the JWT against Okta\'s OIDC discovery endpoint\n'
    '• SAP OData calls: The MCP server passes the token as Authorization: Bearer to SAP\n\n'
    'This is possible because SAP is configured to trust the same Okta authorization server.'
)

# ── 1. Okta Configuration ──
doc.add_heading('1. Okta Configuration', level=1)

doc.add_heading('1.1 Okta Application Setup', level=2)
doc.add_paragraph('Create a Web Application in Okta with the following settings:')
t = doc.add_table(rows=7, cols=2, style='Light Grid Accent 1')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
data = [
    ('Setting', 'Value'),
    ('Application type', 'Web Application'),
    ('Grant type', 'Authorization Code'),
    ('Sign-in redirect URIs', 'http://localhost:8085/callback (Python)\nhttps://oauth.pstmn.io/v1/callback (Postman)'),
    ('Scopes', 'openid email'),
    ('Client authentication', 'Client secret (Basic Auth)'),
    ('State', 'abc'),
]
for i, (k, v) in enumerate(data):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v
    if i == 0:
        for cell in t.rows[i].cells:
            cell.paragraphs[0].runs[0].bold = True

doc.add_heading('1.2 Okta Endpoints', level=2)
doc.add_paragraph('All endpoints use the default custom authorization server (/oauth2/default):')
t = doc.add_table(rows=6, cols=2, style='Light Grid Accent 1')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
data = [
    ('Endpoint', 'URL'),
    ('Authorization', 'https://trial-1053860.okta.com/oauth2/default/v1/authorize'),
    ('Token', 'https://trial-1053860.okta.com/oauth2/default/v1/token'),
    ('OIDC Discovery', 'https://trial-1053860.okta.com/oauth2/default/.well-known/openid-configuration'),
    ('JWKS', 'https://trial-1053860.okta.com/oauth2/default/v1/keys'),
    ('Issuer', 'https://trial-1053860.okta.com/oauth2/default'),
]
for i, (k, v) in enumerate(data):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v
    if i == 0:
        for cell in t.rows[i].cells:
            cell.paragraphs[0].runs[0].bold = True

doc.add_heading('1.3 Access Token Claims', level=2)
t = doc.add_table(rows=5, cols=2, style='Light Grid Accent 1')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
data = [
    ('Claim', 'Value'),
    ('iss', 'https://trial-1053860.okta.com/oauth2/default'),
    ('aud', 'https://trial-1053860.okta.com/oauth2/default'),
    ('cid', '0oa10vth79kZAuXGt698'),
    ('scp', '["openid", "email"]'),
]
for i, (k, v) in enumerate(data):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v
    if i == 0:
        for cell in t.rows[i].cells:
            cell.paragraphs[0].runs[0].bold = True

doc.add_heading('1.4 Important Notes', level=2)
doc.add_paragraph(
    '• The access_token (not the id_token) is used for SAP authentication\n'
    '• The token issuer MUST match what is configured in both SAP and AgentCore\n'
    '• Using /oauth2/v1/ (org auth server) produces tokens with issuer '
    'https://trial-1053860.okta.com — this does NOT match SAP config\n'
    '• Using /oauth2/default/v1/ produces tokens with issuer '
    'https://trial-1053860.okta.com/oauth2/default — this MATCHES SAP config'
)

# ── 2. SAP Configuration ──
doc.add_heading('2. SAP S/4HANA Configuration', level=1)

doc.add_heading('2.1 OIDC Provider Setup', level=2)
doc.add_paragraph('SAP is configured to trust Okta as an OIDC identity provider:')
t = doc.add_table(rows=6, cols=2, style='Light Grid Accent 1')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
data = [
    ('Setting', 'Value'),
    ('Issuer', 'https://trial-1053860.okta.com/oauth2/default'),
    ('Client ID', '0oa10vth79kZAuXGt698'),
    ('User Mapping Claim', 'sub claim (OIDC standard)'),
    ('User Mapping Mechanism', 'E-Mail'),
    ('JWKS Download URL', 'https://trial-1053860.okta.com/oauth2/default/v1/keys'),
]
for i, (k, v) in enumerate(data):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v
    if i == 0:
        for cell in t.rows[i].cells:
            cell.paragraphs[0].runs[0].bold = True

doc.add_heading('2.2 SAP OData Service', level=2)
t = doc.add_table(rows=4, cols=2, style='Light Grid Accent 1')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
data = [
    ('Setting', 'Value'),
    ('SAP Server', 'https://vhcals4hci.awspoc.club'),
    ('OData Service', '/sap/opu/odata/sap/API_SALES_ORDER_SRV/A_SalesOrder'),
    ('Authentication', 'Bearer token (Okta access_token)'),
]
for i, (k, v) in enumerate(data):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v
    if i == 0:
        for cell in t.rows[i].cells:
            cell.paragraphs[0].runs[0].bold = True

doc.add_heading('2.3 SAP Prerequisites', level=2)
doc.add_paragraph(
    '• SAP ICM must have HTTPS enabled\n'
    '• Okta\'s SSL root CA certificate must be imported in STRUST (SSL Client PSE)\n'
    '• SAP must be able to reach Okta endpoints outbound (for JWKS key download)\n'
    '• If SAP shows RFC error "ThSAPOCMINIT, CM_PRODUCT_SPECIFIC_ERROR cmRc=20", '
    'it means SAP cannot reach Okta — check network/proxy settings'
)

# ── 3. AgentCore Configuration ──
doc.add_heading('3. Amazon Bedrock AgentCore Configuration', level=1)

doc.add_heading('3.1 MCP Server Tools', level=2)
doc.add_paragraph('The MCP server (sap_odata_mcp_server.py) exposes two tools:')
t = doc.add_table(rows=3, cols=3, style='Light Grid Accent 1')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
data = [
    ('Tool', 'Description', 'Parameters'),
    ('get_sales_orders', 'Get sales orders with pagination', 'bearer_token, top, skip'),
    ('get_sales_order_by_id', 'Get a specific sales order', 'bearer_token, sales_order_id'),
]
for i, row_data in enumerate(data):
    for j, val in enumerate(row_data):
        t.rows[i].cells[j].text = val
        if i == 0:
            t.rows[i].cells[j].paragraphs[0].runs[0].bold = True

doc.add_paragraph(
    '\nKey implementation details:\n'
    '• Uses FastMCP with stateless_http=True (required for AgentCore Runtime)\n'
    '• Listens on 0.0.0.0:8000/mcp\n'
    '• Makes OData calls using httpx with verify=False (for self-signed SAP certs)\n'
    '• The bearer_token parameter is passed by the client — it is the Okta access_token'
)

doc.add_heading('3.2 Runtime Deployment Configuration', level=2)
t = doc.add_table(rows=7, cols=2, style='Light Grid Accent 1')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
data = [
    ('Setting', 'Value'),
    ('Agent Name', 'sap_odata_mcp_server'),
    ('Protocol', 'MCP'),
    ('Platform', 'linux/arm64'),
    ('Network Mode', 'PUBLIC'),
    ('ECR', 'Auto-created'),
    ('Execution Role', 'Auto-created'),
]
for i, (k, v) in enumerate(data):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v
    if i == 0:
        for cell in t.rows[i].cells:
            cell.paragraphs[0].runs[0].bold = True

doc.add_heading('3.3 Inbound Auth (JWT Authorizer)', level=2)
doc.add_paragraph('AgentCore Runtime is configured with a custom JWT authorizer:')
p = doc.add_paragraph()
run = p.add_run(
    'authorizer_configuration = {\n'
    '    "customJWTAuthorizer": {\n'
    '        "allowedClients": ["0oa10vth79kZAuXGt698"],\n'
    '        "discoveryUrl": "https://trial-1053860.okta.com/oauth2/default/.well-known/openid-configuration"\n'
    '    }\n'
    '}'
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph(
    '\nAgentCore validates incoming requests by:\n'
    '1. Extracting the Bearer token from the Authorization header\n'
    '2. Fetching JWKS keys from the discovery URL\n'
    '3. Validating the JWT signature, expiry, and cid claim against allowedClients'
)

doc.add_heading('3.4 Deployment Steps', level=2)
p = doc.add_paragraph()
run = p.add_run(
    '# Set environment\n'
    'export OKTA_DOMAIN="trial-1053860.okta.com"\n'
    'export OKTA_CLIENT_ID="0oa10vth79kZAuXGt698"\n'
    'export AWS_DEFAULT_REGION="us-east-1"\n\n'
    '# Deploy (requires Docker running)\n'
    'cd sap-agentcore-mcp\n'
    'pip install -r requirements_server.txt\n'
    'python deploy_mcp_server.py'
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph(
    '\nThe deploy script:\n'
    '1. Configures AgentCore Runtime with MCP protocol and Okta JWT authorizer\n'
    '2. Builds a Docker image with the MCP server code\n'
    '3. Pushes to ECR\n'
    '4. Creates/updates the AgentCore Runtime\n'
    '5. Waits for READY status\n'
    '6. Stores the Agent ARN in SSM Parameter Store (/sap_mcp_server/agent_arn)'
)

doc.add_heading('3.5 Server Dependencies', level=2)
p = doc.add_paragraph()
run = p.add_run(
    'mcp>=1.10.0\n'
    'httpx\n'
    'boto3\n'
    'bedrock-agentcore<=0.1.5\n'
    'bedrock-agentcore-starter-toolkit==0.1.14'
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

# ── 4. Python Client ──
doc.add_heading('4. Python Client (sap_agent_client.py)', level=1)

doc.add_heading('4.1 Authentication Flow', level=2)
doc.add_paragraph(
    'The client implements the OAuth2 Authorization Code flow:\n\n'
    '1. Starts a local HTTP server on localhost:8085 for the callback\n'
    '2. Opens the browser to Okta\'s authorize endpoint\n'
    '3. User authenticates with Okta (MFA supported)\n'
    '4. Okta redirects to localhost:8085/callback with an authorization code\n'
    '5. Client exchanges the code for tokens using Basic Auth (client_id:client_secret)\n'
    '6. Uses the access_token for both AgentCore and SAP'
)

doc.add_heading('4.2 MCP Connection', level=2)
doc.add_paragraph('The client connects to AgentCore Runtime via Streamable HTTP:')
p = doc.add_paragraph()
run = p.add_run(
    'mcp_url = f"https://bedrock-agentcore.{region}.amazonaws.com"\n'
    '          f"/runtimes/{encoded_arn}/invocations?qualifier=DEFAULT"\n\n'
    'headers = {\n'
    '    "authorization": f"Bearer {access_token}",\n'
    '    "Content-Type": "application/json",\n'
    '}'
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('4.3 Running the Client', level=2)
p = doc.add_paragraph()
run = p.add_run(
    '# PowerShell\n'
    '$env:OKTA_CLIENT_SECRET="<your-secret>"\n'
    'cd sap-agentcore-mcp\n'
    'pip install -r requirements_client.txt\n'
    'python sap_agent_client.py "show me top 2 sales orders"'
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

# ── 5. Local Testing ──
doc.add_heading('5. Local Testing (test_sap_local.py)', level=1)
doc.add_paragraph(
    'A standalone script that bypasses AgentCore to test the Okta to SAP connection directly. '
    'Useful for debugging token or SAP issues without AgentCore in the middle.\n\n'
    'The script:\n'
    '1. Performs the same Okta Authorization Code flow\n'
    '2. Prints full token contents with decoded claims\n'
    '3. Calls SAP OData directly with the access_token\n'
)
p = doc.add_paragraph()
run = p.add_run(
    '# PowerShell\n'
    '$env:OKTA_CLIENT_SECRET="<your-secret>"\n'
    'python test_sap_local.py'
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

# ── 6. Troubleshooting ──
doc.add_heading('6. Troubleshooting', level=1)

doc.add_heading('401 from AgentCore Runtime', level=2)
doc.add_paragraph(
    'Cause: Token issuer does not match the discovery URL configured in AgentCore.\n'
    'Fix: Ensure the Okta authorize/token URLs use /oauth2/default/v1/ (not /oauth2/v1/).\n'
    'Verify: Decode the access_token and check the iss claim matches the discovery URL issuer.'
)

doc.add_heading('401 from SAP', level=2)
doc.add_paragraph(
    'Cause 1: Wrong token type — SAP expects the access_token, not the id_token.\n'
    'Cause 2: Token issuer does not match SAP OIDC config '
    '(must be https://trial-1053860.okta.com/oauth2/default).\n'
    'Cause 3: SAP cannot reach Okta to download JWKS keys (RFC error cmRc=20).\n'
    'Fix: Use test_sap_local.py to test the token directly against SAP.'
)

doc.add_heading('CodeBuild Failed during deployment', level=2)
doc.add_paragraph(
    'Cause: Dependency version mismatch in requirements_server.txt.\n'
    'Fix: Use pinned versions: bedrock-agentcore<=0.1.5, bedrock-agentcore-starter-toolkit==0.1.14'
)

doc.add_heading('Agent already exists (ConflictException)', level=2)
doc.add_paragraph(
    'Cause: Re-deploying without updating.\n'
    'Fix: Use auto_update_on_conflict=True in agentcore_runtime.launch()'
)

doc.add_heading('Environment variable not persisting', level=2)
doc.add_paragraph(
    'PowerShell: Use $env:OKTA_CLIENT_SECRET="value"\n'
    'CMD: Use set OKTA_CLIENT_SECRET=value\n'
    'Must be set in the same terminal session as the python command.'
)

# ── 7. File Reference ──
doc.add_heading('7. File Reference', level=1)
t = doc.add_table(rows=8, cols=2, style='Light Grid Accent 1')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
data = [
    ('File', 'Purpose'),
    ('sap_odata_mcp_server.py', 'MCP server with SAP OData tools (deployed to AgentCore)'),
    ('deploy_mcp_server.py', 'Deploys MCP server to AgentCore Runtime with Okta auth'),
    ('sap_agent_client.py', 'Python client — Okta login → AgentCore → SAP'),
    ('test_sap_local.py', 'Local test — Okta login → SAP directly (no AgentCore)'),
    ('check_token.py', 'Debug utility — prints token claims'),
    ('requirements_server.txt', 'Server-side dependencies'),
    ('requirements_client.txt', 'Client-side dependencies'),
]
for i, (k, v) in enumerate(data):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v
    if i == 0:
        for cell in t.rows[i].cells:
            cell.paragraphs[0].runs[0].bold = True

# Save
doc.save('SAP_AgentCore_Okta_Integration_Guide.docx')
print("Word document generated: SAP_AgentCore_Okta_Integration_Guide.docx")
